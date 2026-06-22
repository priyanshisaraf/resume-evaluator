import os
from io import BytesIO
import docx
import fitz

async def extract_text_from_file(uploaded_file) -> str:
    filename = uploaded_file.filename or ""
    file_ext = os.path.splitext(filename)[-1].lower()
    file_bytes = await uploaded_file.read()

    if not file_bytes:
        raise ValueError("The uploaded file is empty.")
    if len(file_bytes) > 10 * 1024 * 1024:
        raise ValueError("Resume must be smaller than 10 MB.")

    if file_ext == ".pdf":
        text = extract_text_from_pdf_faster(file_bytes)
    elif file_ext == ".docx":
        text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError("Only PDF and DOCX resumes are supported.")

    if not text.strip():
        raise ValueError("No readable text was found in the resume.")
    return text

def extract_text_from_pdf_faster(file_bytes: bytes) -> str:
    with fitz.open(stream=file_bytes, filetype="pdf") as document:
        return "\n".join(page.get_text() for page in document)

def extract_text_from_docx(file_bytes: bytes) -> str:
    document = docx.Document(BytesIO(file_bytes))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + table_cells)

def truncate_text(text: str, max_words: int = 1000) -> str:
    words = text.split()
    return ' '.join(words[:max_words])
